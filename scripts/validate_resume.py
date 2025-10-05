from docx import Document
from docx.shared import Pt
import sys
import re

def validate_resume(resume_path):
    """
    Validate the generated resume for:
    1. No remaining placeholders
    2. Correct formatting (Times New Roman, 10pt)
    3. Complete content replacement

    Returns:
        tuple: (is_valid, errors_list)
    """
    errors = []

    try:
        doc = Document(resume_path)
    except Exception as e:
        return False, [f"ERROR: Could not open document: {e}"]

    # Check for remaining placeholders
    placeholder_patterns = [
        r'\[.*?\]',  # Any text in brackets
        r'W\d-B\d',  # Work bullet identifiers (W1-B1, W2-B1, etc.)
        r'L-B\d',    # Leadership bullet identifiers (L-B1, L-B2, etc.)
    ]

    # Check tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, para in enumerate(cell.paragraphs):
                    text = para.text.strip()

                    # Check for placeholders
                    for pattern in placeholder_patterns:
                        matches = re.findall(pattern, text)
                        if matches:
                            errors.append(f"PLACEHOLDER FOUND in Table {table_idx}, Row {row_idx}, Cell {cell_idx}: {matches}")

                    # Check formatting for non-empty paragraphs
                    if text and len(para.runs) > 0:
                        for run_idx, run in enumerate(para.runs):
                            if run.text.strip():  # Only check non-empty runs
                                # Check font name
                                if run.font.name and run.font.name != 'Times New Roman':
                                    errors.append(f"WRONG FONT in Table {table_idx}, Row {row_idx}, Cell {cell_idx}: '{run.font.name}' (expected 'Times New Roman')")

                                # Check font size (10pt for body text)
                                if run.font.size:
                                    size_pt = run.font.size.pt
                                    # Allow some common sizes: 10pt for body, 11pt for headers, 12pt/14pt/22pt for name/titles
                                    if size_pt not in [10.0, 11.0, 12.0, 14.0, 22.0]:
                                        errors.append(f"WRONG SIZE in Table {table_idx}, Row {row_idx}, Cell {cell_idx}: {size_pt}pt (expected 10pt for body text)")

    # Check regular paragraphs
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # Check for placeholders
        for pattern in placeholder_patterns:
            matches = re.findall(pattern, text)
            if matches:
                errors.append(f"PLACEHOLDER FOUND in Paragraph {para_idx}: {matches}")

        # Check formatting
        if text and len(para.runs) > 0:
            for run_idx, run in enumerate(para.runs):
                if run.text.strip():
                    if run.font.name and run.font.name != 'Times New Roman':
                        errors.append(f"WRONG FONT in Paragraph {para_idx}: '{run.font.name}' (expected 'Times New Roman')")

                    if run.font.size:
                        size_pt = run.font.size.pt
                        if size_pt not in [10.0, 11.0, 12.0, 14.0, 22.0]:
                            errors.append(f"WRONG SIZE in Paragraph {para_idx}: {size_pt}pt")

    # Determine if valid
    is_valid = len(errors) == 0

    return is_valid, errors

if __name__ == "__main__":
    # Get resume path from command line or use default
    if len(sys.argv) > 1:
        resume_path = sys.argv[1]
    else:
        print("Usage: python validate_resume.py <path_to_resume.docx>")
        sys.exit(1)

    print("=" * 80)
    print(f"VALIDATING RESUME: {resume_path}")
    print("=" * 80)

    is_valid, errors = validate_resume(resume_path)

    if is_valid:
        print("\n[PASS] VALIDATION PASSED - Resume is correctly formatted!")
        print("\nNo issues found:")
        print("  - All placeholders replaced")
        print("  - Correct font (Times New Roman)")
        print("  - Correct sizing (10pt body text)")
    else:
        print(f"\n[FAIL] VALIDATION FAILED - Found {len(errors)} issue(s):")
        print()
        for i, error in enumerate(errors, 1):
            print(f"{i}. {error}")
        print("\n" + "=" * 80)
        print("Please review and fix these issues.")

    print("\n" + "=" * 80)

    sys.exit(0 if is_valid else 1)
