# scripts/utils/docx.py
"""Manipulation de documents Word - remplacement de placeholders."""

import re
from docx import Document
from docx.shared import Pt
from .config import FONT_NAME, FONT_SIZE_PT


def replace_placeholder(paragraph, placeholder: str, replacement: str, support_markdown: bool = False) -> bool:
    """
    Remplace un placeholder dans un paragraphe en préservant le formatting.

    Args:
        paragraph: Paragraphe python-docx
        placeholder: Texte à remplacer
        replacement: Texte de remplacement
        support_markdown: Si True, parse **bold** syntax

    Returns:
        True si remplacement effectué, False sinon
    """
    full_text = paragraph.text

    if placeholder not in full_text:
        return False

    # Récupérer le formatting original
    original_font_name = FONT_NAME
    original_font_size = Pt(FONT_SIZE_PT)

    for run in paragraph.runs:
        if run.text.strip():
            if run.font.name:
                original_font_name = run.font.name
            if run.font.size:
                original_font_size = run.font.size
            break

    # Remplacer le texte
    new_text = full_text.replace(placeholder, replacement)

    # Appliquer le nouveau texte
    if support_markdown and '**' in new_text:
        _apply_markdown_formatting(paragraph, new_text, original_font_name, original_font_size)
    else:
        _apply_simple_text(paragraph, new_text, original_font_name, original_font_size)

    return True


def _apply_simple_text(paragraph, text: str, font_name: str, font_size):
    """Applique du texte simple sans markdown."""
    # Clear all runs
    for run in paragraph.runs:
        run.text = ""

    # Add text to first run
    if paragraph.runs:
        first_run = paragraph.runs[0]
    else:
        first_run = paragraph.add_run()

    first_run.text = text
    first_run.font.name = font_name
    first_run.font.size = font_size
    first_run.font.bold = None
    first_run.font.italic = None


def _apply_markdown_formatting(paragraph, text: str, font_name: str, font_size):
    """Parse **bold** syntax et applique le formatting Word."""
    # Clear all runs
    for run in paragraph.runs:
        run.text = ""

    # Split by **text** pattern
    parts = re.split(r'(\*\*.*?\*\*)', text)

    for part in parts:
        if not part:
            continue

        if part.startswith('**') and part.endswith('**'):
            # Bold text
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.font.name = font_name
            run.font.size = font_size
            run.font.bold = True
        else:
            # Normal text
            run = paragraph.add_run(part)
            run.font.name = font_name
            run.font.size = font_size
            run.font.bold = None


def replace_all_placeholders(doc: Document, replacements: dict, support_markdown: bool = False):
    """
    Remplace tous les placeholders dans un document.

    Args:
        doc: Document python-docx
        replacements: Dict {placeholder: replacement}
        support_markdown: Si True, parse **bold** syntax
    """
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, replacement in replacements.items():
                        if replacement:  # Skip empty replacements
                            replace_placeholder(paragraph, placeholder, replacement, support_markdown)

    # Process regular paragraphs
    for paragraph in doc.paragraphs:
        for placeholder, replacement in replacements.items():
            if replacement:
                replace_placeholder(paragraph, placeholder, replacement, support_markdown)


def remove_empty_paragraphs(doc: Document, placeholders_to_remove: list[str] = None):
    """
    Supprime les paragraphes vides ou contenant certains placeholders.

    Args:
        doc: Document python-docx
        placeholders_to_remove: Liste de placeholders dont les paragraphes doivent être supprimés
    """
    if placeholders_to_remove:
        # Remove paragraphs containing specific placeholders
        paragraphs_to_remove = []
        for paragraph in doc.paragraphs:
            for placeholder in placeholders_to_remove:
                if placeholder in paragraph.text:
                    paragraphs_to_remove.append(paragraph)
                    break

        for paragraph in paragraphs_to_remove:
            p_element = paragraph._element
            p_element.getparent().remove(p_element)

    # Remove trailing empty paragraphs
    while len(doc.paragraphs) > 0:
        last_para = doc.paragraphs[-1]
        if not last_para.text.strip():
            p_element = last_para._element
            p_element.getparent().remove(p_element)
        else:
            break
