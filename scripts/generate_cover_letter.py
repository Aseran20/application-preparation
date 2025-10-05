from docx import Document
from docx.shared import Pt
import os
from datetime import datetime
import shutil
import re
import subprocess

def parse_markdown_and_apply_formatting(paragraph, text, original_font_name, original_font_size):
    """
    Parse markdown syntax (**bold**) and apply Word formatting.
    Handles text like: **Title** — rest of text
    """
    # Clear all runs
    for run in paragraph.runs:
        run.text = ""

    # Pattern to match **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)

    for part in parts:
        if not part:
            continue

        if part.startswith('**') and part.endswith('**'):
            # Bold text
            bold_text = part[2:-2]  # Remove **
            run = paragraph.add_run(bold_text)
            run.font.name = original_font_name
            run.font.size = original_font_size
            run.font.bold = True
        else:
            # Normal text
            run = paragraph.add_run(part)
            run.font.name = original_font_name
            run.font.size = original_font_size
            run.font.bold = None

def replace_paragraph_text_cl(paragraph, placeholder, replacement_text):
    """
    Replace placeholder in paragraph while preserving ALL formatting.
    This keeps the exact style, alignment, fonts from the template.
    """
    # Get full text of paragraph
    full_text = paragraph.text

    # Check if placeholder exists
    if placeholder not in full_text:
        return False

    # Store original formatting from first run with text
    original_font_name = None
    original_font_size = None

    for run in paragraph.runs:
        if run.text.strip():
            original_font_name = run.font.name
            original_font_size = run.font.size
            break

    # Replace text
    new_text = full_text.replace(placeholder, replacement_text)

    # Check if replacement text contains markdown
    if '**' in new_text:
        # Parse markdown and apply formatting
        parse_markdown_and_apply_formatting(paragraph, new_text, original_font_name, original_font_size)
    else:
        # Simple replacement without markdown
        # Clear all runs
        for run in paragraph.runs:
            run.text = ""

        # Add the first run back with replacement
        if paragraph.runs:
            first_run = paragraph.runs[0]
        else:
            first_run = paragraph.add_run()

        # Apply original formatting
        if original_font_name:
            first_run.font.name = original_font_name
        if original_font_size:
            first_run.font.size = original_font_size

        first_run.text = new_text

    return True

def generate_cover_letter(
    company_name,
    recipient_name,
    street_number,
    postal_city_country,
    intro_paragraph,
    body_paragraph_1,
    body_paragraph_2,
    body_paragraph_3,
    additional_context,
    company_attraction,
    closing_paragraph,
    output_folder
):
    """
    Generate customized cover letter from template v2.

    Args:
        company_name: Company name
        recipient_name: Hiring manager name or "Hiring Manager"
        street_number: Street and number
        postal_city_country: Postal code, city, country
        intro_paragraph: Opening paragraph
        body_paragraph_1: First numbered reason (1.)
        body_paragraph_2: Second numbered reason (2.)
        body_paragraph_3: Third numbered reason (3.)
        additional_context: Additional context paragraph (optional)
        company_attraction: What attracts you to the company (optional)
        closing_paragraph: Closing paragraph
        output_folder: Folder to save the cover letter

    Returns:
        tuple: (folder_name, output_path)
    """

    # Load template
    doc = Document("templates/Cover Letter - Adrian Turion.docx")

    # Get current date
    current_date = datetime.now().strftime("%B %d, %Y")

    # Salutation
    salutation = f"Dear {recipient_name},"

    # Placeholders to replace
    placeholders = {
        "[DATE]": current_date,
        "[RECIPIENT_NAME]": recipient_name,
        "[STREET_NUMBER]": street_number,
        "[POSTAL_CITY_COUNTRY]": postal_city_country,
        "[COMPANY_NAME]": company_name,
        "[SALUTATION]": salutation,
        "[INTRO_PARAGRAPH]": intro_paragraph,
        "[BODY_PARAGRAPH_1]": body_paragraph_1,
        "[BODY_PARAGRAPH_2]": body_paragraph_2,
        "[BODY_PARAGRAPH_3]": body_paragraph_3,
        "[ADDITIONAL_CONTEXT]": additional_context if additional_context else "",
        "[COMPANY_ATTRACTION]": company_attraction if company_attraction else "",
        "[CLOSING_PARAGRAPH]": closing_paragraph,
    }

    # First pass: identify paragraphs to remove
    paragraphs_to_remove = []
    for i, paragraph in enumerate(doc.paragraphs):
        for placeholder, replacement in placeholders.items():
            if placeholder in paragraph.text:
                # Mark empty optional paragraphs for removal
                if not replacement and placeholder in ["[ADDITIONAL_CONTEXT]", "[COMPANY_ATTRACTION]"]:
                    paragraphs_to_remove.append(paragraph)

    # Remove empty optional paragraphs
    for paragraph in paragraphs_to_remove:
        p_element = paragraph._element
        p_element.getparent().remove(p_element)

    # Second pass: replace placeholders
    for paragraph in doc.paragraphs:
        for placeholder, replacement in placeholders.items():
            if placeholder in paragraph.text and replacement:
                replace_paragraph_text_cl(paragraph, placeholder, replacement)

    # Save cover letter
    output_path = os.path.join(output_folder, "Cover_Letter_Adrian_Turion.docx")
    doc.save(output_path)

    # Create PDF folder and convert to PDF
    pdf_folder = os.path.join(output_folder, "PDF")
    os.makedirs(pdf_folder, exist_ok=True)

    pdf_output_path = os.path.join(pdf_folder, "Cover_Letter_Adrian_Turion.pdf")

    try:
        # Convert DOCX to PDF using LibreOffice (cross-platform)
        abs_output_path = os.path.abspath(output_path)
        abs_pdf_folder = os.path.abspath(pdf_folder)

        # Try LibreOffice conversion
        subprocess.run([
            "soffice", "--headless", "--convert-to", "pdf",
            "--outdir", abs_pdf_folder, abs_output_path
        ], check=True, capture_output=True)

        print(f"[PDF] {pdf_output_path}")
    except (subprocess.CalledProcessError, FileNotFoundError):
        # LibreOffice not available, try docx2pdf (Windows)
        try:
            from docx2pdf import convert
            convert(abs_output_path, pdf_output_path)
            print(f"[PDF] {pdf_output_path}")
        except ImportError:
            print("[WARNING] PDF conversion skipped (install LibreOffice or docx2pdf)")
            pdf_output_path = None

    return output_folder, output_path, pdf_output_path

# Example usage (for testing)
if __name__ == "__main__":
    # Test with Glencore example
    company_name = "Glencore"
    recipient_name = "Hiring Manager"
    company_address = ""  # Not provided

    intro = """I am writing to express my strong interest in the Commercial Graduate Programme at Glencore. The opportunity to develop commercial aptitude and data analysis skills within your dynamic trading environment directly aligns with my experience in fast-paced M&A deal-making and my passion for applying analytical rigor to complex business challenges."""

    body1 = """During my time at Auraïa Capital Advisory, I managed seven concurrent mandates ranging from EUR 5M to EUR 50M, developing the resilience and multitasking abilities essential for trading support roles. Working in a lean three-person team gave me individual responsibility from day one—coordinating with clients, advisors, and counterparties while maintaining accuracy under tight deadlines. To increase efficiency, I built an AI-powered tool that reduced buyer sourcing time by 90%, demonstrating my ability to leverage technology for operational excellence."""

    body2 = """What particularly draws me to Glencore is your entrepreneurial culture and commitment to individual ownership. As the founder of two ventures—including an AI software adopted by advisory firms—I bring the entrepreneurial abilities and initiative valued in your programme. Your emphasis on accuracy and strong communication resonates with my experience presenting to C-suite executives and managing client relationships throughout complex transaction lifecycles."""

    body3 = ""  # Not needed for this one

    closing = """I am excited about the opportunity to contribute to Glencore's commercial teams and develop my career in commodities trading. I would welcome the chance to discuss how my analytical skills, entrepreneurial mindset, and proven resilience can add value to your programme."""

    folder = "jobs/Glencore_Commercial_Graduate_05-10-2025"

    folder_name, output_path = generate_cover_letter(
        company_name, recipient_name, company_address,
        intro, body1, body2, body3, closing, folder
    )

    print(f"Cover letter generated successfully!")
    print(f"Saved to: {output_path}")
