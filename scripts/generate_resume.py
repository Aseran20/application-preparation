from docx import Document
from docx.shared import Pt
import os
from datetime import datetime
import shutil
import subprocess

def replace_paragraph_text(paragraph, placeholder, replacement_text):
    """
    Replace placeholder in paragraph while preserving formatting.
    Handles cases where placeholder is split across multiple runs.
    """
    # Get full text of paragraph
    full_text = paragraph.text

    # Check if placeholder exists
    if placeholder not in full_text:
        return False

    # Clear all runs
    for run in paragraph.runs:
        run.text = ""

    # Add the first run back with replacement
    if paragraph.runs:
        first_run = paragraph.runs[0]
    else:
        first_run = paragraph.add_run()

    # Preserve formatting
    first_run.font.name = 'Times New Roman'
    first_run.font.size = Pt(10)
    first_run.font.bold = None
    first_run.font.italic = None

    # Replace text
    new_text = full_text.replace(placeholder, replacement_text)
    first_run.text = new_text

    return True

def generate_resume(professional_summary, auraia_bullets, rc_group_bullet,
                   europ_assistance_bullet, leadership_bullets, courses,
                   skills, company, position):
    """
    Generate customized resume from template.

    Args:
        professional_summary: 2-3 line intro text
        auraia_bullets: List of 3 accomplishments
        rc_group_bullet: Single accomplishment (can be longer)
        europ_assistance_bullet: Single accomplishment (can be longer)
        leadership_bullets: List of 3 leadership experiences
        courses: Comma-separated string of 5 courses
        skills: One-line string of technical skills (may include | separator for subsections)
        company: Company name for folder
        position: Position name for folder
    """

    # Load template
    doc = Document("templates/Resume - Adrian Turion.docx")

    # Placeholders to replace
    placeholders = {
        "[Your professional summary here - describe your background, expertise, and career goals]": professional_summary,
        "[Describe your key responsibility or achievement here W1-B1]": auraia_bullets[0],
        "[Describe your key responsibility or achievement here W1-B2]": auraia_bullets[1],
        "[Describe your key responsibility or achievement here W1-B3]": auraia_bullets[2],
        "[Describe your key responsibility or achievement here W2-B1]": rc_group_bullet,
        "[Describe your key responsibility or achievement here W3-B1]": europ_assistance_bullet,
        "[Describe your key responsibility or achievement here L-B1]": leadership_bullets[0],
        "[Describe your key responsibility or achievement here L-B2]": leadership_bullets[1],
        "[Describe your key responsibility or achievement here L-B3]": leadership_bullets[2],
        "[Relevant coursework here]": courses,
        "[Relevant software here]": skills,
    }

    # Process all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Try to replace each placeholder
                    for placeholder, replacement in placeholders.items():
                        replace_paragraph_text(paragraph, placeholder, replacement)

    # Process regular paragraphs (just in case)
    for paragraph in doc.paragraphs:
        for placeholder, replacement in placeholders.items():
            replace_paragraph_text(paragraph, placeholder, replacement)

    # Create output folder
    date_str = datetime.now().strftime("%d-%m-%Y")
    folder_name = f"jobs/{company}_{position}_{date_str}"
    os.makedirs(folder_name, exist_ok=True)

    # Remove trailing empty paragraphs to avoid blank pages in PDF
    while len(doc.paragraphs) > 0:
        last_para = doc.paragraphs[-1]
        # Check if last paragraph is empty or only whitespace
        if not last_para.text.strip():
            # Remove the paragraph
            p_element = last_para._element
            p_element.getparent().remove(p_element)
        else:
            break

    # Save customized resume
    output_path = os.path.join(folder_name, "Resume_Adrian_Turion.docx")
    doc.save(output_path)

    # Copy job description
    shutil.copy("job_description.md", os.path.join(folder_name, "job_description.md"))

    # Create PDF folder and convert to PDF
    pdf_folder = os.path.join(folder_name, "PDF")
    os.makedirs(pdf_folder, exist_ok=True)

    pdf_output_path = os.path.join(pdf_folder, "Resume_Adrian_Turion.pdf")
    temp_pdf_path = os.path.join(pdf_folder, "Resume_Adrian_Turion_temp.pdf")

    try:
        # Convert DOCX to PDF using LibreOffice (cross-platform)
        abs_output_path = os.path.abspath(output_path)
        abs_pdf_folder = os.path.abspath(pdf_folder)

        # Try LibreOffice conversion
        subprocess.run([
            "soffice", "--headless", "--convert-to", "pdf",
            "--outdir", abs_pdf_folder, abs_output_path
        ], check=True, capture_output=True)

        # Keep only first page of PDF using PyPDF2
        try:
            from PyPDF2 import PdfReader, PdfWriter

            reader = PdfReader(pdf_output_path)
            writer = PdfWriter()

            # Add only the first page
            if len(reader.pages) > 0:
                writer.add_page(reader.pages[0])

            # Write to temp file then replace original
            with open(temp_pdf_path, 'wb') as output_file:
                writer.write(output_file)

            # Replace original with single-page version
            shutil.move(temp_pdf_path, pdf_output_path)
            print(f"[PDF] {pdf_output_path} (1 page)")
        except ImportError:
            print(f"[PDF] {pdf_output_path} (PyPDF2 not installed - may contain blank pages)")
        except Exception as e:
            print(f"[PDF] {pdf_output_path} (page trimming failed: {e})")

    except (subprocess.CalledProcessError, FileNotFoundError):
        # LibreOffice not available, try docx2pdf (Windows)
        try:
            from docx2pdf import convert
            convert(abs_output_path, pdf_output_path)

            # Keep only first page of PDF using PyPDF2
            try:
                from PyPDF2 import PdfReader, PdfWriter

                reader = PdfReader(pdf_output_path)
                writer = PdfWriter()

                # Add only the first page
                if len(reader.pages) > 0:
                    writer.add_page(reader.pages[0])

                # Write to temp file then replace original
                with open(temp_pdf_path, 'wb') as output_file:
                    writer.write(output_file)

                # Replace original with single-page version
                shutil.move(temp_pdf_path, pdf_output_path)
                print(f"[PDF] {pdf_output_path} (1 page)")
            except ImportError:
                print(f"[PDF] {pdf_output_path} (PyPDF2 not installed - may contain blank pages)")
            except Exception as e:
                print(f"[PDF] {pdf_output_path} (page trimming failed: {e})")

        except ImportError:
            print("[WARNING] PDF conversion skipped (install LibreOffice or docx2pdf)")
            pdf_output_path = None

    return folder_name, output_path, pdf_output_path

# This allows the script to be imported and used by the /create command
if __name__ == "__main__":
    # Example usage - will be replaced by /create command
    professional_summary = "Finance graduate with hands-on M&A experience across healthcare, technology, and life sciences sectors. Proven ability to thrive in fast-paced environments, combining strong analytical and data analysis skills with entrepreneurial mindset. Seeking to leverage commercial aptitude and cross-border deal exposure in Glencore's dynamic trading environment."

    auraia_bullets = [
        "Staffed on 7 live mandates (sell-side, buy-side, debt advisory) across healthcare, tech, and life sciences sectors with deal values ranging from EUR 5M to EUR 50M",
        "End-to-end execution exposure in a lean 3-person deal team (sourcing, valuation, due diligence, negotiations, closing)",
        "Conducted comprehensive industry research and identified strategic buyers for cross-border deals spanning Europe, US, and Asia"
    ]

    rc_group_bullet = "Worked directly with the CEO of the Healthcare and Retail divisions on evaluating multiple acquisition targets in France and Belgium, preparing comprehensive investment memos including market analysis, competitive positioning, and financial projections for 3 potential acquisitions"

    europ_assistance_bullet = "Built automated financial dashboards in Excel and PowerBI to track KPIs and performed variance analysis on monthly P&L results, supporting the CFO and Head of Controlling on strategic projects for the Swiss subsidiary (EUR 50M+ revenues)"

    leadership_bullets = [
        "Launched a startup that customized Nike sneakers using upcycled materials, managing product development, supplier relationships, and social media marketing to 500+ followers",
        "Drove business development at AIESEC, a volunteering student organization, securing 3 corporate partnerships and increasing volunteer placements by 40%",
        "Co-organized TEDx HEC Lausanne with 200+ attendees, managing speaker relations, logistics, and a €15,000 budget"
    ]

    courses = "Risk Management, Derivatives, Applied Econometrics and Statistics, Data Science for Finance, Quantitative Asset and Risk Management"

    skills = "Financial Modeling: Excel (Advanced), Bloomberg Terminal, Capital IQ, PowerBI | Data Science: Python, SQL, R"

    company = "Glencore"
    position = "Commercial_Graduate"

    folder_name, output_path = generate_resume(
        professional_summary, auraia_bullets, rc_group_bullet,
        europ_assistance_bullet, leadership_bullets, courses,
        skills, company, position
    )

    print(f"Resume generated successfully!")
    print(f"Output folder: {folder_name}")
    print(f"Files created:")
    print(f"   - Resume_Adrian_Turion.docx")
    print(f"   - job_description.md")
