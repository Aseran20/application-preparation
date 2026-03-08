# scripts/utils/docx.py
"""Manipulation de documents Word - remplacement de placeholders."""

import re
from docx import Document
from docx.shared import Pt
from .config import FONT_NAME, FONT_SIZE_PT


def replace_placeholder(paragraph, placeholder: str, replacement: str, support_markdown: bool = False) -> bool:
    """
    Remplace un placeholder dans un paragraphe en préservant le formatting de chaque run.

    Fonctionne même quand le placeholder est éclaté sur plusieurs runs (comportement
    courant de Word). Le texte de remplacement hérite du formatting du run où commence
    le placeholder ; les autres runs restent intacts.

    Args:
        paragraph: Paragraphe python-docx
        placeholder: Texte à remplacer
        replacement: Texte de remplacement
        support_markdown: ignoré (conservé pour compatibilité)

    Returns:
        True si remplacement effectué, False sinon
    """
    full_text = paragraph.text

    if placeholder not in full_text:
        return False

    ph_start = full_text.index(placeholder)
    ph_end = ph_start + len(placeholder)

    # Map each run to its character range in the full text
    pos = 0
    run_ranges = []
    for run in paragraph.runs:
        rlen = len(run.text)
        run_ranges.append((pos, pos + rlen, run))
        pos += rlen

    replacement_inserted = False

    for r_start, r_end, run in run_ranges:
        # No overlap with placeholder
        if r_end <= ph_start or r_start >= ph_end:
            continue

        # Portion of this run's text that falls inside the placeholder
        overlap_start = max(r_start, ph_start) - r_start
        overlap_end = min(r_end, ph_end) - r_start

        before = run.text[:overlap_start]
        after = run.text[overlap_end:]

        if not replacement_inserted:
            run.text = before + replacement + after
            replacement_inserted = True
        else:
            run.text = before + after  # remove placeholder portion only

    return True


def replace_placeholder_with_newlines(doc, paragraph, placeholder: str, replacement: str, support_markdown: bool = False) -> bool:
    """
    Remplace un placeholder et crée de nouveaux paragraphes pour chaque \\n.

    Args:
        doc: Document python-docx
        paragraph: Paragraphe contenant le placeholder
        placeholder: Texte à remplacer
        replacement: Texte de remplacement (peut contenir \\n)
        support_markdown: ignoré (conservé pour compatibilité)

    Returns:
        True si remplacement effectué, False sinon
    """
    full_text = paragraph.text

    if placeholder not in full_text:
        return False

    # Si pas de newline, utiliser la méthode simple
    if '\n' not in replacement:
        return replace_placeholder(paragraph, placeholder, replacement)

    # Récupérer le formatting original du premier run non-vide
    original_font_name = FONT_NAME
    original_font_size = Pt(FONT_SIZE_PT)

    for run in paragraph.runs:
        if run.text.strip():
            if run.font.name:
                original_font_name = run.font.name
            if run.font.size:
                original_font_size = run.font.size
            break

    # Récupérer le style du paragraphe
    paragraph_style = paragraph.style
    paragraph_format = paragraph.paragraph_format

    # Split le replacement par newlines
    parts = replacement.split('\n')

    # Premier paragraphe : remplacer avec la première partie
    replace_placeholder(paragraph, placeholder, parts[0])

    # Trouver l'index du paragraphe dans le document
    p_element = paragraph._element
    parent = p_element.getparent()

    # Créer les paragraphes suivants
    current_element = p_element
    for part in parts[1:]:
        if not part.strip():  # Skip empty parts
            continue

        new_para = doc.add_paragraph()
        new_para.style = paragraph_style

        if paragraph_format.space_after:
            new_para.paragraph_format.space_after = paragraph_format.space_after
        if paragraph_format.space_before:
            new_para.paragraph_format.space_before = paragraph_format.space_before

        run = new_para.add_run(part)
        run.font.name = original_font_name
        run.font.size = original_font_size

        new_p_element = new_para._element
        parent.remove(new_p_element)
        current_element.addnext(new_p_element)
        current_element = new_p_element

    return True


def replace_all_placeholders(doc: Document, replacements: dict, support_markdown: bool = False):
    """
    Remplace tous les placeholders dans un document.

    Args:
        doc: Document python-docx
        replacements: Dict {placeholder: replacement}
        support_markdown: ignoré (conservé pour compatibilité)
    """
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, replacement in replacements.items():
                        if replacement:  # Skip empty replacements
                            replace_placeholder(paragraph, placeholder, replacement)

    # Process regular paragraphs
    for paragraph in doc.paragraphs:
        for placeholder, replacement in replacements.items():
            if replacement:
                replace_placeholder(paragraph, placeholder, replacement)


def remove_empty_paragraphs(doc: Document, placeholders_to_remove: list[str] = None):
    """
    Supprime les paragraphes vides ou contenant certains placeholders.

    Args:
        doc: Document python-docx
        placeholders_to_remove: Liste de placeholders dont les paragraphes doivent être supprimés
    """
    if placeholders_to_remove:
        paragraphs_to_remove = []
        for paragraph in doc.paragraphs:
            for placeholder in placeholders_to_remove:
                if placeholder in paragraph.text:
                    paragraphs_to_remove.append(paragraph)
                    break

        for paragraph in paragraphs_to_remove:
            p_element = paragraph._element
            p_element.getparent().remove(p_element)

    # Remove trailing empty paragraphs
    while len(doc.paragraphs) > 0:
        last_para = doc.paragraphs[-1]
        if not last_para.text.strip():
            p_element = last_para._element
            p_element.getparent().remove(p_element)
        else:
            break
